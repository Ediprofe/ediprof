#!/usr/bin/env python3
"""
md-to-docx.py - Convierte Markdown a Word con control total sobre imágenes

Uso: python3 md-to-docx.py input1.md [input2.md ...] -o output.docx [-t template.docx]
"""

import sys
import os
import re
import subprocess
import tempfile
import hashlib
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def parse_args():
    """Parsea argumentos de línea de comandos."""
    args = sys.argv[1:]
    input_files = []
    output_file = "output.docx"
    template_file = None
    
    i = 0
    while i < len(args):
        if args[i] in ('-o', '--output'):
            output_file = args[i + 1]
            i += 2
        elif args[i] in ('-t', '--template'):
            template_file = args[i + 1]
            i += 2
        else:
            input_files.append(args[i])
            i += 1
    
    return input_files, output_file, template_file

def get_project_root():
    """Obtiene la raíz del proyecto."""
    script_dir = Path(__file__).parent
    return script_dir.parent

def convert_svg_to_png(svg_path, temp_dir, project_root):
    """Convierte SVG a PNG usando Playwright."""
    svg_path = Path(svg_path)
    if not svg_path.exists():
        return None
    
    # Nombre único para el PNG
    hash_str = hashlib.md5(str(svg_path).encode()).hexdigest()[:8]
    png_name = f"{svg_path.stem}_{hash_str}.png"
    png_path = Path(temp_dir) / png_name
    
    # Usar el script de Node para la conversión
    node_script = project_root / "scripts" / "svg-to-png.mjs"
    if node_script.exists():
        try:
            subprocess.run(
                ["node", str(node_script), str(svg_path), str(png_path), "1.5"],
                capture_output=True, check=True
            )
            if png_path.exists():
                return png_path
        except subprocess.CalledProcessError:
            pass
    
    return None

def add_paragraph_from_text(doc, text, style=None):
    """Agrega un párrafo con el texto dado."""
    p = doc.add_paragraph(text, style=style)
    return p

def process_markdown_line(doc, line, temp_dir, project_root, in_code_block):
    """Procesa una línea de markdown y la agrega al documento."""
    
    # Detectar bloques de código
    if line.strip().startswith('```'):
        return not in_code_block, True
    
    if in_code_block:
        # Agregar como código
        p = doc.add_paragraph(line)
        p.style = 'No Spacing'
        run = p.runs[0] if p.runs else p.add_run()
        run.font.name = 'Consolas'
        run.font.size = Pt(10)
        return in_code_block, True
    
    # Líneas vacías
    if not line.strip():
        return in_code_block, False
    
    # Headers
    if line.startswith('# '):
        doc.add_heading(line[2:].strip(), level=1)
        return in_code_block, True
    if line.startswith('## '):
        doc.add_heading(line[3:].strip(), level=2)
        return in_code_block, True
    if line.startswith('### '):
        doc.add_heading(line[4:].strip(), level=3)
        return in_code_block, True
    if line.startswith('#### '):
        doc.add_heading(line[5:].strip(), level=4)
        return in_code_block, True
    
    # Imágenes: ![alt](/images/...)
    img_match = re.match(r'!\[([^\]]*)\]\((/images/[^)]+)\)', line)
    if img_match:
        alt_text = img_match.group(1)
        img_path = img_match.group(2)
        
        # Resolver ruta absoluta
        if img_path.endswith('.svg'):
            abs_path = project_root / "public" / img_path.lstrip('/')
            png_path = convert_svg_to_png(abs_path, temp_dir, project_root)
            if png_path:
                # Agregar imagen centrada, al ancho de la página
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                # Ancho de 6 pulgadas (aproximadamente el ancho útil de la página)
                run.add_picture(str(png_path), width=Inches(6))
                return in_code_block, True
        else:
            # PNG/JPG existente
            abs_path = project_root / "public" / img_path.lstrip('/')
            if abs_path.exists():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(str(abs_path), width=Inches(6))
                return in_code_block, True
        
        return in_code_block, False
    
    # Listas
    if line.strip().startswith('- ') or line.strip().startswith('* '):
        text = line.strip()[2:]
        doc.add_paragraph(text, style='List Bullet')
        return in_code_block, True
    
    # Listas numeradas
    num_match = re.match(r'^\d+\.\s+(.+)$', line.strip())
    if num_match:
        doc.add_paragraph(num_match.group(1), style='List Number')
        return in_code_block, True
    
    # Separador horizontal
    if line.strip() in ('---', '***', '___'):
        # Agregar línea horizontal
        p = doc.add_paragraph('─' * 50)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return in_code_block, True
    
    # Texto normal (procesar negritas e itálicas básicas)
    text = line.strip()
    # Por ahora, agregar como texto plano
    if text:
        p = doc.add_paragraph(text)
        return in_code_block, True
    
    return in_code_block, False

def convert_md_to_docx(input_files, output_file, template_file=None):
    """Convierte archivos Markdown a un documento Word."""
    
    project_root = get_project_root()
    
    # Crear documento (con o sin plantilla)
    if template_file and Path(template_file).exists():
        doc = Document(template_file)
    else:
        doc = Document()
    
    # Directorio temporal para imágenes convertidas
    with tempfile.TemporaryDirectory() as temp_dir:
        
        for idx, input_file in enumerate(input_files):
            print(f"  📄 Procesando: {Path(input_file).name}")
            
            with open(input_file, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Quitar frontmatter YAML
            if content.startswith('---'):
                end_fm = content.find('---', 3)
                if end_fm != -1:
                    content = content[end_fm + 3:].strip()
            
            in_code_block = False
            
            for line in content.split('\n'):
                in_code_block, _ = process_markdown_line(
                    doc, line, temp_dir, project_root, in_code_block
                )
            
            # Agregar salto de página entre archivos
            if idx < len(input_files) - 1:
                doc.add_page_break()
        
        # Guardar documento
        doc.save(output_file)
        print(f"✅ Documento creado: {output_file}")

def main():
    print("━" * 50)
    
    input_files, output_file, template_file = parse_args()
    
    if not input_files:
        print("❌ Error: No se especificaron archivos de entrada")
        sys.exit(1)
    
    print(f"📚 Exportando {len(input_files)} lecciones...")
    convert_md_to_docx(input_files, output_file, template_file)

if __name__ == "__main__":
    main()
